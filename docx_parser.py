import re
import io
from docx import Document

ROLE_NORMALISE: dict[str, str | None] = {
    "focal teacher":             "focal_teacher",
    "focal teacher (as)":        "focal_teacher",
    "non-focal teacher":         "non_focal_teacher",
    "non focal teacher":         "non_focal_teacher",
    "co-teacher":                "non_focal_teacher",
    "co teacher":                "non_focal_teacher",
    "research team member":      "research_team",
    "research team":             "research_team",
    "researcher":                "research_team",
    "student":                   "student",
    "unknown student speaking":  "unknown_student",
    "unknown student":           "unknown_student",
    "unknown":                   "unknown_student",
    "inaudible":                 None,   
}

SPEAKER_RE   = re.compile(r'^([A-Z]{1,3}):\s*(.*)', re.DOTALL)
TIMESTAMP_RE = re.compile(r'@\s*(\d{1,2}:\d{2})')
KEY_ENTRY_RE = re.compile(r'^([A-Z]{1,3})\s*=\s*(.+)', re.IGNORECASE)
STAGE_RE     = re.compile(r'^\[.*\]$')

_MONTH_MAP = {
    "january": "01", "february": "02", "march": "03", "april": "04",
    "may": "05",     "june": "06",     "july": "07",  "august": "08",
    "september": "09", "october": "10", "november": "11", "december": "12",
    "jan": "01", "feb": "02", "mar": "03", "apr": "04",
    "jun": "06", "jul": "07", "aug": "08", "sep": "09",
    "oct": "10", "nov": "11", "dec": "12",
}


def _normalise_role(raw: str) -> str | None:
    key = raw.strip().lower().rstrip(".")
    if key in ROLE_NORMALISE:
        return ROLE_NORMALISE[key]
    for pattern, role in ROLE_NORMALISE.items():
        if key.startswith(pattern):
            return role
    return "unknown"


def _parse_date(raw: str) -> str:
    """Convert '30 September 2025' or '2025-09-30' → 'YYYY-MM-DD'."""
    raw = raw.strip()
    m = re.match(r"(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})", raw)
    if m:
        day = m.group(1).zfill(2)
        mon = _MONTH_MAP.get(m.group(2).lower(), "00")
        return f"{m.group(3)}-{mon}-{day}"
    m2 = re.match(r"(\d{4})-(\d{2})-(\d{2})", raw)
    if m2:
        return m2.group(0)
    return "unknown"


def _is_translation(text: str) -> bool:
    if not text or len(text.strip()) < 3:
        return False
    if SPEAKER_RE.match(text.strip()):
        return False
    if TIMESTAMP_RE.match(text.strip()):
        return False
    if STAGE_RE.match(text.strip()):
        return False
    return True



def parse_key_section(doc: Document) -> tuple[dict[str, str | None], dict]:
    speaker_roles: dict[str, str | None] = {}
    lesson_meta: dict = {}

    in_key = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text.lower() == "key:":
            in_key = True
            continue

        if in_key:
            m = KEY_ENTRY_RE.match(text)
            if m:
                code = m.group(1).upper()
                role = _normalise_role(m.group(2))
                speaker_roles[code] = role
            else:
                dm = re.search(r"Date of Lesson:\s*(.+?)(?:\n|Lesson|$)", text, re.IGNORECASE)
                if dm:
                    lesson_meta["lesson_date"] = _parse_date(dm.group(1))

                lm = re.search(r"Lesson Length:\s*(\d+:\d+)", text, re.IGNORECASE)
                if lm:
                    lesson_meta["lesson_length"] = lm.group(1)

                sm = re.search(r"Class starts at:\s*@?\s*(\d+:\d+)", text, re.IGNORECASE)
                if sm:
                    lesson_meta["class_start"] = sm.group(1)

                if text.startswith("@"):
                    break

    return speaker_roles, lesson_meta


def parse_docx(
    docx_bytes: bytes,
    transcript_id: str,
    lesson_date: str,            # overridden if found in document
) -> list[dict]:
    doc = Document(io.BytesIO(docx_bytes))

   
    speaker_roles, lesson_meta = parse_key_section(doc)


    effective_date = lesson_meta.get("lesson_date") or lesson_date

    if not speaker_roles:
        raise ValueError(
            f"No speaker roles found in Key section of '{transcript_id}'. "
            "Check that the document starts with a 'Key:' block."
        )

    paragraphs = [p.text.strip() for p in doc.paragraphs]

    turns: list[dict] = []
    current_timestamp = lesson_meta.get("class_start", "00:00")
    i = 0

    while i < len(paragraphs):
        text = paragraphs[i]

        if not text:
            i += 1
            continue

      
        ts_match = TIMESTAMP_RE.search(text)
        if ts_match and not SPEAKER_RE.match(text):
            current_timestamp = ts_match.group(1)
            i += 1
            continue

        speaker_match = SPEAKER_RE.match(text)
        if speaker_match:
            code    = speaker_match.group(1).strip()
            content = speaker_match.group(2).strip()
            role    = speaker_roles.get(code)

           
            if role is None:
                i += 1
                continue

           
            translation = None
            j = i + 1
            while j < len(paragraphs) and not paragraphs[j].strip():
                j += 1
            if j < len(paragraphs) and _is_translation(paragraphs[j]):
                translation = paragraphs[j].strip()
                i = j  

            turns.append({
                "speaker_code":  code,
                "speaker_role":  role,
                "timestamp":     current_timestamp,
                "text":          content,
                "translation":   translation,
                "transcript_id": transcript_id,
                "lesson_date":   effective_date,
                "has_inaudible": "[INA]" in content,
            })

        i += 1

    return turns
